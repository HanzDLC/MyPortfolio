"""
Generate an ATS-friendly Resume .docx (and PDF via Word) that mirrors the
on-site resume styling (templates/resume_document.html + .cv-* CSS):
Times New Roman serif, bold-label contact grid, dark-blue section rules.

AUTO-FITS TO ONE PAGE: tries progressively tighter configs and stops at the
first that renders a single page (memory: feedback-resume-one-page).

Run:  python generate_resume_docx.py
Outputs: Hanz_de_la_Cruz_Resume.docx  and  Hanz_de_la_Cruz_Resume.pdf
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = "Hanz_de_la_Cruz_Resume.docx"
OUT_PDF = "Hanz_de_la_Cruz_Resume.pdf"

# Resume content — overridable by variant scripts (e.g. generate_resume_va.py).
# Defaults below produce the canonical resume verbatim.
SUMMARY = ("Full-stack developer with hands-on experience building production-grade web "
           "applications and autonomous AI systems. Skilled in modern JavaScript frameworks, "
           "Python, database design, and AI/LLM integration. Known for shipping reliable, "
           "well-architected systems that solve real operational problems and deliver "
           "measurable impact.")
SKILLS = [
    ("Technical:",
     "Full-Stack Development (Next.js, React, FastAPI, Node.js), Frontend (TypeScript, "
     "Tailwind CSS, HTML/CSS/JavaScript), Backend (FastAPI, Python, PostgreSQL, Prisma ORM), "
     "Agentic AI and LLM Integration (Claude, ChatGPT, Paperclip), Prompt Engineering, "
     "Workflow Automation (n8n, OpenClaw), QA Testing and Bug Documentation, Data Analysis "
     "and Visualization"),
    ("Professional:",
     "Communication, Team Collaboration, Analytical Problem Solving, Time Management, "
     "Attention to Detail, Ownership and Coachability, Adaptability"),
]

SERIF = "Times New Roman"
INK = RGBColor(0x1A, 0x1A, 0x1A)      # #1a1a1a name/title/date/labels
GRAY3 = RGBColor(0x33, 0x33, 0x33)    # #333 contact values
GRAY4 = RGBColor(0x44, 0x44, 0x44)    # #444 detail/institution
LINKBLUE = RGBColor(0x1A, 0x52, 0x76) # #1a5276 portfolio + section rule
RULE_HEX = "1A5276"

# Most-generous → tightest. First config that yields 1 page wins.
CONFIGS = [
    dict(base=10.5, line=1.05, sec=6, mt=0.30, ms=0.6),
    dict(base=10.0, line=1.02, sec=5, mt=0.28, ms=0.6),
    dict(base=10.0, line=0.98, sec=4, mt=0.25, ms=0.55),
    dict(base=9.5,  line=0.98, sec=4, mt=0.25, ms=0.55),
    dict(base=9.5,  line=0.94, sec=3, mt=0.22, ms=0.5),
    dict(base=9.0,  line=0.96, sec=3, mt=0.22, ms=0.5),
    dict(base=9.0,  line=0.90, sec=2, mt=0.20, ms=0.45),
]


def _set_font(run, size, bold=False, italic=False, color=INK, name=SERIF):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def build(cfg):
    BASE = cfg["base"]
    RIGHT_TAB = Inches(8.5 - 2 * cfg["ms"])
    LABEL_TAB = Inches(0.72)  # contact value column (min-width 65px label)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(cfg["mt"])
        s.bottom_margin = Inches(cfg["mt"])
        s.left_margin = Inches(cfg["ms"])
        s.right_margin = Inches(cfg["ms"])

    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(BASE)
    nrpr = normal.element.get_or_add_rPr()
    nrf = nrpr.find(qn("w:rFonts"))
    if nrf is None:
        nrf = OxmlElement("w:rFonts"); nrpr.append(nrf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        nrf.set(qn(attr), SERIF)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = cfg["line"]

    def name_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        _set_font(r, BASE + 8, bold=True, color=INK)
        spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "10")  # letter-spacing 0.5px
        r._element.get_or_add_rPr().append(spc)

    def contact_row(label, value, link=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(LABEL_TAB, WD_TAB_ALIGNMENT.LEFT)
        rl = p.add_run(label)
        _set_font(rl, BASE, bold=True, color=INK)
        p.add_run("\t")
        rv = p.add_run(value)
        if link:
            _set_font(rv, BASE, color=LINKBLUE)
            rv.font.underline = True
        else:
            _set_font(rv, BASE, color=GRAY3)

    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(cfg["sec"])
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        _set_font(r, BASE + 4, bold=True, color=INK)
        ppr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")        # ~1.5pt (CSS 2px)
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), RULE_HEX)
        pbdr.append(bottom)
        ppr.append(pbdr)

    def role_header(title, date, subtitle=None, institution=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        rt = p.add_run(title)
        _set_font(rt, BASE + 1, bold=True, color=INK)
        p.add_run("\t")
        rd = p.add_run(date)
        _set_font(rd, BASE, bold=True, color=INK)
        if subtitle or institution:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
            rs = p2.add_run(subtitle or "")
            _set_font(rs, BASE, italic=True, color=GRAY4)
            if institution:
                p2.add_run("\t")
                ri = p2.add_run(institution)
                _set_font(ri, BASE - 0.5, color=GRAY4)

    def bullets(items):
        for it in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.28)
            if isinstance(it, tuple):
                lead, rest = it
                rb = p.add_run(lead); _set_font(rb, BASE, bold=True, color=INK)
                rr = p.add_run(rest); _set_font(rr, BASE, color=GRAY3)
            else:
                r = p.add_run(it); _set_font(r, BASE, color=GRAY3)

    def para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _set_font(p.add_run(text), BASE, color=GRAY3)

    def labeled(label, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(label + " "); _set_font(rb, BASE, bold=True, color=INK)
        rr = p.add_run(text); _set_font(rr, BASE, color=GRAY3)

    # ---- content (mirrors templates/resume_document.html) ----
    name_heading("Hanz Uriel A. de la Cruz")
    contact_row("Address:", "Landheights Ville, Tagbak, Jaro, Iloilo City")
    contact_row("Phone:", "0945-332-4472 / 0968-697-3611")
    contact_row("Email:", "hdlcruz03@gmail.com")
    contact_row("Portfolio:", "https://hanzdlc-portfolio.vercel.app/", link=True)

    section_title("Professional Summary")
    para(SUMMARY)

    section_title("Relevant Experience")

    role_header("Chief Technology Officer (CTO) and Forward Deployed Engineer",
                "2026 - Present",
                "SMAPS — School Management AI-Powered System", "Iloilo, Philippines")
    bullets([
        "Lead technical direction, architecture, and the product roadmap for SMAPS, a web-based school management information system with DepEd-aligned electronic class record (ECR) grading.",
        ("Built the core platform: ", "school, subject, teacher, and grading-configuration "
         "modules plus an integrated ECR grading system, using Base44, MongoDB, and n8n "
         "automation accelerated by prompt engineering."),
        "Replaced manual, fragmented school workflows with accounting-transparency and audit-trail processes across the system.",
    ])

    role_header("AI Specialist — Confidential (US Commercial Flooring Company)",
                "May 2026 - Present",
                "Lead-Response and Proposal Automations", "Remote")
    bullets([
        "Building lead-response and proposal automations for a US commercial flooring company",
        ("Built an AI proposal tool", " that auto-fills project fields from pasted lead notes, "
         "computes a live-editable estimate, and auto-drafts the proposal with .xlsx/.docx export "
         "(Next.js, FastAPI, Claude)"),
        ("Built an AI lead-intelligence pipeline", " that monitors public news and permit feeds, "
         "deduplicates them into enriched project records, and emails a daily opportunity digest "
         "(FastAPI, Supabase, Claude, Docker)"),
        ("Built a login-gated internal systems hub", " with per-system SOPs and developer "
         "documentation (React, FastAPI, PostgreSQL, Nginx); automating admin workflows and "
         "knowledge systems with Claude Cowork"),
        "Role details confidential per NDA — available upon request with written authorization",
    ])

    role_header("Lead Full-Stack Developer and AI Engineer",
                "March 2026 - Present",
                "Zilla Media — ARIA, DriveXP, OpenClaw and Hermes Agentic AI", "Remote")
    bullets([
        ("ARIA: ", "Led development of a full-stack AI marketing platform with 6 autonomous "
                   "agents — delivered ~3x faster CEO chat, hardened Paperclip sub-agent "
                   "delegation, and self-healing Claude CLI (Next.js 14, FastAPI, Supabase, "
                   "Paperclip AI, Socket.IO)."),
        ("DriveXP: ", "Built a car rental platform with a 16-page admin dashboard and 9 "
                      "interconnected database models (Next.js 16, React 19, TypeScript, "
                      "Prisma 7, PostgreSQL)."),
        ("OpenClaw: ", "Architected and deployed an autonomous workflow orchestration system "
                       "that handles task automation and real-time monitoring without manual "
                       "oversight."),
        ("Hermes Agent (Nous Research): ", "Deployed a self-improving agent hub with persistent "
                                           "memory, cron automations, and multi-platform "
                                           "messaging (Telegram/Discord/Slack) — migrated from "
                                           "OpenClaw."),
    ])

    section_title("Certifications")
    bullets([
        "Google AI Essentials — Coursera (2026)",
        "Data Science Essentials with Python — Cisco Networking Academy",
        "Python Essentials 1 — Cisco Networking Academy",
        "Apply AI: Analyze Customer Reviews — Cisco Networking Academy",
        "Introduction to Cybersecurity — Cisco Networking Academy",
        "DevOps Basics — DICT Philippines",
        "Civil Service Examination Passer — Professional Level (2024)",
    ])

    section_title("Education")
    role_header("Bachelor of Science in Information Technology",
                "August 2022 - Present",
                "College of Computing and Informatics",
                "Iloilo Science and Technology University")

    section_title("Core Skills")
    for _label, _text in SKILLS:
        labeled(_label, _text)

    doc.save(OUT_DOCX)


def to_pdf(word):
    d = word.Documents.Open(os.path.abspath(OUT_DOCX))
    d.SaveAs(os.path.abspath(OUT_PDF), FileFormat=17)
    d.Close()


def page_count():
    import fitz
    doc = fitz.open(OUT_PDF)
    n = doc.page_count
    doc.close()
    return n


def fit_and_save():
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    chosen = None
    try:
        for cfg in CONFIGS:
            build(cfg)
            to_pdf(word)
            n = page_count()
            print(f"  base={cfg['base']} line={cfg['line']} -> {n} page(s)")
            chosen = cfg
            if n == 1:
                break
    finally:
        word.Quit()
    return chosen


if __name__ == "__main__":
    cfg = fit_and_save()
    print(f"Saved {OUT_DOCX} + {OUT_PDF}  |  Times New Roman {cfg['base']}pt  "
          f"|  {os.path.getsize(OUT_PDF)/1048576:.3f} MB")
